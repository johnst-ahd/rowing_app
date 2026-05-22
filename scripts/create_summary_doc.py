"""Generate Word summary for rowing PWA + Traccar integration."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT = Path(__file__).resolve().parents[1] / "docs" / "Rowing_App_PWA_Summary.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Rowing Sensor PWA — Structure & Limitations", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Summary for integrating a phone-based Progressive Web App (PWA) with the "
        "existing traccar-overlay project and Traccar GPS backend."
    )
    meta = doc.add_paragraph()
    meta.add_run("Repository: ").bold = True
    meta.add_run("https://github.com/JohnSt-AHD/rowing_app\n")
    meta.add_run("Existing host project: ").bold = True
    meta.add_run("traccar-overlay (maps, live race, RowSafe)\n")
    meta.add_run("Date: ").bold = True
    meta.add_run("May 2026")

    add_heading(doc, "1. Purpose", 1)
    doc.add_paragraph(
        "Record heart rate (HR), GPS, and accelerometer data on athletes’ phones "
        "during rowing sessions, and send that data into the existing RNZ tracking "
        "stack so traccar-overlay can display live position and related telemetry."
    )

    add_heading(doc, "2. Recommended architecture", 1)
    doc.add_paragraph(
        "Treat the solution as two products: rowing_app (phone recorder) and "
        "traccar-overlay (display and operations). Do not merge them into one repo."
    )

    add_heading(doc, "2.1 Data flow", 2)
    add_bullet(doc, " — Phone PWA captures sensors, buffers offline in IndexedDB, uploads when online.")
    add_bullet(doc, " — GPS positions sent to Traccar via OsmAnd HTTP protocol (port 5055). Extra fields (e.g. HR) can be attached as position attributes.")
    add_bullet(doc, " — HR and high-rate accelerometer sent to a small ingest API (Vercel), stored as time-series batches keyed by device/session.")
    add_bullet(doc, " — traccar-overlay continues to poll Traccar for positions; optionally extended later to read ingest API for HR/accel.")

    add_heading(doc, "2.2 Suggested repository structure (rowing_app)", 2)
    structure = doc.add_paragraph()
    structure.style = "No Spacing"
    run = structure.add_run(
        "rowing_app/\n"
        "  apps/recorder-pwa/     — Vite + TypeScript PWA (manifest, service worker)\n"
        "    src/sensors/         — gps, motion, heart-rate (Web Bluetooth)\n"
        "    src/session/         — recorder, IndexedDB outbox\n"
        "    src/upload/          — traccar (OsmAnd), telemetry-api\n"
        "  api/                   — Vercel serverless ingest (optional phase 2)\n"
        "  packages/telemetry-types/ — shared JSON/TypeScript schemas\n"
        "  docs/                  — architecture and integration notes\n"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    add_heading(doc, "2.3 Deployment", 2)
    add_bullet(doc, " — Vercel: static PWA + api/ingest")
    add_bullet(doc, " — traccar-overlay stays on existing Vercel deployment")
    add_bullet(doc, " — Traccar admin credentials remain server-side only; phones use device IDs and session tokens")

    add_heading(doc, "3. GPS update rates and comparison", 1)

    add_heading(doc, "3.1 Three separate layers", 2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Typical rate"
    hdr[2].text = "Notes"
    rows = [
        ("Phone → Traccar", "0.3–1 Hz (aggressive Client)", "OS and app settings dominate"),
        ("Traccar database", "Same as ingest", "Optional filter.minPeriod can drop fixes"),
        ("traccar-overlay map UI", "0.03–0.33 Hz (3–10 s poll)", "User setting; default 10 s"),
    ]
    for i, (a, b, c) in enumerate(rows, 1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c

    add_heading(doc, "3.2 Traccar Client (native app)", 2)
    add_bullet(doc, "Interval and Distance use OR logic — reports when either threshold is met.")
    add_bullet(doc, "Fastest interval caps maximum send rate; Highest accuracy keeps GPS on for better consistency.")
    add_bullet(doc, "Practical live tracking: about 1 Hz (e.g. Interval 1–3 s, low distance).")
    add_bullet(doc, "Fleet/safety style: every 30–60 s or every 50–100 m.")
    add_bullet(doc, "Strong background tracking on Android (wake lock); iOS still subject to OS limits.")

    add_heading(doc, "3.3 PWA (browser)", 2)
    add_bullet(doc, "Geolocation watchPosition has no fixed Hz; browser may rate-limit and uses “significant change” logic.")
    add_bullet(doc, "Foreground (screen on, active tab): roughly 0.2–2 Hz, device-dependent.")
    add_bullet(doc, "Background / locked screen: often sparse or no updates — especially on iOS.")
    add_bullet(doc, "Generally underperforms Traccar Client for continuous track density.")

    add_heading(doc, "3.4 traccar-overlay display", 2)
    add_bullet(doc, "Map refresh configurable from 3 s to 60 s (default 10 s) — independent of GPS ingest rate.")
    add_bullet(doc, "Even 1 Hz stored in Traccar may appear smoother only if overlay poll is set to 3–5 s during live races.")

    add_heading(doc, "4. Sensor limitations", 1)

    add_heading(doc, "4.1 GPS", 2)
    doc.add_paragraph(
        "PWA GPS is suitable for active recording sessions with the app in the foreground. "
        "For all-day or pocket tracking, the official Traccar Client app is more reliable."
    )

    add_heading(doc, "4.2 Accelerometer", 2)
    doc.add_paragraph(
        "DeviceMotion API works over HTTPS; iOS may require explicit permission. "
        "Sample at 10–25 Hz for general motion; 50 Hz only if needed for stroke analysis, "
        "with batched uploads via the ingest API (not Traccar positions)."
    )

    add_heading(doc, "4.3 Heart rate", 2)
    doc.add_paragraph(
        "Web Bluetooth (GATT heart_rate service) works well on Android Chrome. "
        "iOS Safari does not support Web Bluetooth natively; alternatives include Android devices, "
        "BLE-capable browsers, or third-party iOS extensions/polyfills (extra user setup). "
        "Connection must be started from a user tap (button), not on page load."
    )

    add_heading(doc, "5. Implementation phases", 1)
    phases = [
        "Scaffold Vite PWA, manifest, service worker, IndexedDB buffer.",
        "GPS only → Traccar OsmAnd; verify on traccar-overlay live map.",
        "Add accelerometer + ingest API.",
        "Add heart rate (Android first; document iOS path).",
        "Session UX: start/stop, reconnect BLE, export for debugging.",
    ]
    for i, phase in enumerate(phases, 1):
        doc.add_paragraph(f"{i}. {phase}", style="List Number")

    add_heading(doc, "6. Practical recommendation for RNZ rowing", 1)
    add_bullet(doc, " — Use Traccar Client or foreground PWA for GPS at 1–3 s intervals.")
    add_bullet(doc, " — Set traccar-overlay map refresh to 3–5 s during live events.")
    add_bullet(doc, " — Route HR and accelerometer through a dedicated ingest API, not high-rate Traccar positions.")

    add_heading(doc, "7. References", 1)
    refs = [
        "GitHub repository: https://github.com/JohnSt-AHD/rowing_app",
        "Traccar OsmAnd protocol: https://www.traccar.org/osmand/",
        "Traccar Client configuration: https://www.traccar.org/client-configuration/",
        "Traccar server filters: https://www.traccar.org/configuration-file/",
        "W3C Geolocation API: https://www.w3.org/TR/geolocation/",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
